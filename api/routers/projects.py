from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from typing import List, Optional
import tempfile
import os
import json as json_module
import zipfile
import base64
import logging

from api.models.schemas import (
    ProjectCreate, ProjectUpdate, ProjectResponse, ProjectFile, FeedbackRequest
)
from api.services.firestore import db_service
from api.services.gemini import gemini_service
from api.services.latex import latex_service
from api.routers.auth import get_current_user

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/projects", tags=["Projects"])

@router.get("", response_model=List[ProjectResponse])
async def list_projects(user: dict = Depends(get_current_user)):
    projects = await db_service.get_user_projects(user["uid"])
    return [_format_project(p) for p in projects]

@router.get("/{project_id}", response_model=ProjectResponse)
async def get_project(project_id: str, user: dict = Depends(get_current_user)):
    project = await db_service.get_project(project_id, user["uid"])
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return _format_project(project)

@router.post("", response_model=ProjectResponse)
async def create_project(request: ProjectCreate, user: dict = Depends(get_current_user)):
    # Get template content
    templates = latex_service.get_sample_templates()
    template_content = templates.get(request.theme, templates["report"])
    
    files = [
        {"name": "main.tex", "content": template_content, "type": "tex"},
        {"name": "references.bib", "content": "@misc{example,\n  author = {Author},\n  title = {Title},\n  year = {2024}\n}", "type": "bib"}
    ]
    
    project = await db_service.create_project(
        uid=user["uid"],
        name=request.name,
        theme=request.theme,
        files=files,
        main_file="main.tex",
        custom_theme=request.custom_theme
    )
    
    return _format_project(project)

@router.post("/save-project")
async def save_project(request: ProjectUpdate, user: dict = Depends(get_current_user)):
    from datetime import datetime, timezone
    files = [f.dict() for f in request.files]
    success = await db_service.update_project(request.project_id, user["uid"], files)

    if not success:
        raise HTTPException(status_code=404, detail="Project not found")

    updated_at = datetime.now(timezone.utc).isoformat()
    return {"message": "Project saved", "project_id": request.project_id, "updated_at": updated_at}

@router.patch("/{project_id}/rename")
async def rename_project(project_id: str, request: dict, user: dict = Depends(get_current_user)):
    name = request.get("name")
    if not name or not name.strip():
        raise HTTPException(status_code=400, detail="Name is required")
    
    success = await db_service.update_project_name(project_id, user["uid"], name.strip())
    
    if not success:
        raise HTTPException(status_code=404, detail="Project not found")
    
    return {"message": "Project renamed", "name": name.strip()}

@router.post("/duplicate-project/{project_id}")
async def duplicate_project(project_id: str, user: dict = Depends(get_current_user)):
    project = await db_service.duplicate_project(project_id, user["uid"])
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    return _format_project(project)

@router.delete("/delete-project/{project_id}")
async def delete_project(project_id: str, user: dict = Depends(get_current_user)):
    success = await db_service.delete_project(project_id, user["uid"])
    if not success:
        raise HTTPException(status_code=404, detail="Project not found")

    return {"message": "Project deleted"}

@router.post("/{project_id}/add-images")
async def add_images_to_project(
    project_id: str,
    images: List[UploadFile] = File(...),
    user: dict = Depends(get_current_user)
):
    """Add one or more image files to an existing project."""
    project = await db_service.get_project(project_id, user["uid"])
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    ext_type_map = {"png": "png", "jpg": "jpg", "jpeg": "jpg", "gif": "png", "bmp": "png", "pdf": "pdf"}
    mime_map = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                "gif": "image/gif", "bmp": "image/bmp", "pdf": "application/pdf"}
    existing_files = project.get("files", [])
    existing_names = {f["name"] for f in existing_files}

    from api.services.storage import upload_image as gcs_upload

    new_files = list(existing_files)
    added = []
    for upload in images:
        raw = await upload.read()
        if not raw:
            continue
        ext = (upload.filename or "image.png").rsplit(".", 1)[-1].lower()
        file_type = ext_type_map.get(ext, "png")
        mime = mime_map.get(ext, "image/png")
        name = upload.filename or f"image.{ext}"
        # Deduplicate name
        base_name = name.rsplit(".", 1)[0]
        candidate = name
        counter = 1
        while candidate in existing_names:
            candidate = f"{base_name}_{counter}.{ext}"
            counter += 1
        # Upload to GCS instead of storing base64 in Firestore
        gcs_ref = gcs_upload(raw, project_id, candidate, mime)
        new_files.append({"name": candidate, "content": gcs_ref, "type": file_type})
        existing_names.add(candidate)
        added.append(candidate)

    await db_service.update_project(project_id, user["uid"], new_files)
    return {"added": added, "project_id": project_id}

def _format_project(project: dict) -> ProjectResponse:
    files = project.get("files", [])
    formatted_files = [
        ProjectFile(
            name=f.get("name", "unknown"),
            content=f.get("content", ""),
            type=f.get("type", "tex")
        ) for f in files
    ]
    
    return ProjectResponse(
        id=project["id"],
        name=project.get("name", "Untitled"),
        files=formatted_files,
        main_file=project.get("main_file", "main.tex"),
        theme=project.get("theme", "report"),
        custom_theme=project.get("custom_theme"),
        created_at=project.get("created_at"),
        updated_at=project.get("updated_at")
    )

# Upload endpoint at root level
upload_router = APIRouter(tags=["Upload"])

@upload_router.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    theme: str = Form("report"),
    custom_theme: Optional[str] = Form(None),
    custom_prompt: Optional[str] = Form(None),
    custom_cls: Optional[str] = Form(None),
    custom_preamble: Optional[str] = Form(None),
    images: Optional[str] = Form(None),  # JSON array of base64 images
    max_tokens: Optional[int] = Form(None),
    user: dict = Depends(get_current_user)
):
    if not file.filename.endswith(('.docx', '.doc', '.pdf')):
        raise HTTPException(status_code=400, detail="Only DOCX, DOC, or PDF files supported")

    # Parse images from JSON
    image_list = None
    if images:
        try:
            image_list = json_module.loads(images)
        except Exception:
            pass

    # Save file temporarily
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, file.filename)

    try:
        content = await file.read()
        with open(file_path, "wb") as f:
            f.write(content)

        # Extract text and figures from document
        extracted_images: list = image_list or []
        try:
            if file.filename.endswith('.pdf'):
                import fitz
                pdf_doc = fitz.open(file_path)
                text_content = ""
                for page in pdf_doc:
                    text_content += page.get_text()
                pdf_doc.close()
            else:
                # Extract text from DOCX (paragraphs + tables)
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                text_parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(
                            cell.text.strip() for cell in row.cells if cell.text.strip()
                        )
                        if row_text:
                            text_parts.append(row_text)
                text_content = "\n".join(text_parts)

                # Extract embedded images from docx (it's a ZIP archive)
                if not extracted_images:
                    mime_map = {
                        "png": "image/png",
                        "jpg": "image/jpeg",
                        "jpeg": "image/jpeg",
                        "gif": "image/gif",
                        "bmp": "image/bmp",
                    }
                    try:
                        with zipfile.ZipFile(file_path, "r") as zf:
                            media_files = [
                                n for n in zf.namelist()
                                if n.startswith("word/media/")
                                and n.split(".")[-1].lower() in mime_map
                            ]
                            for media_name in media_files[:5]:  # max 5 figures
                                ext = media_name.split(".")[-1].lower()
                                data = zf.read(media_name)
                                b64 = base64.b64encode(data).decode("utf-8")
                                mime = mime_map.get(ext, "image/jpeg")
                                extracted_images.append(f"data:{mime};base64,{b64}")
                    except Exception as img_err:
                        logger.warning(f"Figure extraction failed: {img_err}")
        except Exception as extract_err:
            logger.warning(f"Document extraction failed, falling back to raw text: {extract_err}")
            text_content = content.decode("utf-8", errors="ignore")[:5000]

        logger.info(
            f"Starting generation: file={file.filename!r}, "
            f"text_len={len(text_content)}, images={len(extracted_images)}, "
            f"has_cls={bool(custom_cls and custom_cls.strip())}"
        )

        # Convert to LaTeX using Gemini
        try:
            latex_content, tokens = await gemini_service.generate_document(
                text_content,
                theme,
                custom_theme,
                custom_prompt=custom_prompt,
                custom_preamble=custom_preamble,
                images=extracted_images or None,
                custom_cls_content=custom_cls if custom_cls and custom_cls.strip() else None,
                max_tokens=max_tokens if max_tokens else 65536,
            )
        except Exception as gen_err:
            logger.error(f"Gemini generation failed ({type(gen_err).__name__}): {gen_err}")
            raise HTTPException(
                status_code=502,
                detail=f"AI generation failed ({type(gen_err).__name__}): {str(gen_err)[:300]}"
            )

        # Update user tokens
        await db_service.update_user_tokens(user["uid"], pro_tokens=tokens)

        # Generate bibliography from \cite{} keys found in the LaTeX output
        bib_content = ""
        try:
            bib_content = await gemini_service._generate_bibliography(text_content, latex_content, None)
            if bib_content:
                logger.info(f"Bibliography generated: {len(bib_content)} chars")
        except Exception as bib_err:
            logger.warning(f"Bibliography generation failed: {bib_err}")

        # Build project files
        project_files = [
            {"name": "main.tex", "content": latex_content, "type": "tex"},
            {"name": "references.bib", "content": bib_content, "type": "bib"},
        ]

        # Add embedded images from the source document as project files (via GCS)
        img_ext_map = {"image/png": "png", "image/jpeg": "jpg", "image/gif": "gif", "image/bmp": "bmp"}
        embedded_image_names: set = set()
        # We need a temporary project ID for GCS paths — use a placeholder; real ID assigned after create
        # Instead, collect image bytes and store after project creation
        pending_images: list = []
        for idx, img_data in enumerate(extracted_images):
            if not img_data.startswith("data:"):
                continue
            mime = img_data.split(";")[0][5:]
            ext = img_ext_map.get(mime, "png")
            img_name = f"figure{idx + 1}.{ext}"
            b64_data = img_data.split(",", 1)[1] if "," in img_data else img_data
            import base64 as _base64
            pending_images.append({"name": img_name, "mime": mime, "ext": ext,
                                    "data": _base64.b64decode(b64_data)})
            embedded_image_names.add(img_name)

        # Add custom class file if provided
        if custom_cls and custom_cls.strip():
            project_files.append({"name": "custom.cls", "content": custom_cls, "type": "cls"})

        project = await db_service.create_project(
            uid=user["uid"],
            name=file.filename.rsplit(".", 1)[0],
            theme=theme,
            files=project_files,
            main_file="main.tex",
            custom_theme=custom_theme,
        )

        project_id = project["id"]

        # Upload embedded images to GCS and add as project files
        if pending_images:
            from api.services.storage import upload_image as gcs_upload
            updated_files = list(project_files)
            for img in pending_images:
                try:
                    gcs_ref = gcs_upload(img["data"], project_id, img["name"], img["mime"])
                    updated_files.append({"name": img["name"], "content": gcs_ref, "type": img["ext"]})
                except Exception as img_err:
                    logger.warning(f"Failed to upload embedded image {img['name']}: {img_err}")
            if len(updated_files) > len(project_files):
                await db_service.update_project(project_id, user["uid"], updated_files)

        # Detect image filenames referenced in LaTeX that are not in the project
        image_refs = gemini_service._extract_image_references(latex_content)
        missing_images = [
            ref for ref in image_refs
            if ref not in embedded_image_names
        ]

        return {
            "project_id": project_id,
            "tokens_used": tokens,
            "missing_images": missing_images,
        }

    finally:
        try:
            os.remove(file_path)
            os.rmdir(temp_dir)
        except Exception:
            pass

# Feedback endpoint
feedback_router = APIRouter(tags=["Feedback"])

@feedback_router.post("/feedback")
async def submit_feedback(
    request: FeedbackRequest,
    user: dict = Depends(get_current_user)
):
    await db_service.save_feedback(request.feedback, user.get("uid"))
    return {"message": "Feedback submitted"}
